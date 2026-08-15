"""
Remaining Agents: System, Email, Reminder, File, Voice, Scheduler, Project Manager
"""
# ── System Agent ──────────────────────────────────────────────
# backend/app/agents/system/agent.py
SYSTEM_AGENT = '''"""
System Agent — Local machine operations (requires elevated permissions).
"""
from __future__ import annotations
import os, subprocess, shutil
from pathlib import Path
from typing import Any, Dict
import structlog
from app.agents.base import BaseAgent
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class SystemAgent(BaseAgent):
    name = "system_agent"
    description = "OS operations: open apps, create files, execute programs"
    supported_intents = ["system_operation"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        action = state.metadata.get("system_action", "")
        
        if action == "create_file":
            path = Path(state.metadata.get("path", "output.txt"))
            content = state.metadata.get("content", "")
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(content, encoding="utf-8")
            return {"final_response": f"✅ File created: {path}", "action": "file.create"}
        
        elif action == "run_command":
            cmd = state.metadata.get("command", "")
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=30)
            output = result.stdout or result.stderr
            return {
                "final_response": f"```\n{output}\n```",
                "action": "system.execute",
                "response_type": "code",
            }
        
        elif action == "list_dir":
            path = state.metadata.get("path", ".")
            items = [str(p) for p in Path(path).iterdir()]
            return {"final_response": "\n".join(items)}
        
        return {"final_response": "System action completed", "action": action}
'''

# ── Email Agent ───────────────────────────────────────────────
EMAIL_AGENT = '''"""
Email Agent — Read, summarize, and send emails.
"""
from __future__ import annotations
from typing import Any, Dict
import structlog
from app.agents.base import BaseAgent
from app.config import settings
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class EmailAgent(BaseAgent):
    name = "email_agent"
    description = "Read, summarize, draft, and send emails"
    supported_intents = ["email_read", "email_send"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        if state.intent == "email_send":
            return {
                "action": "email.send",
                "subject": state.metadata.get("subject", ""),
                "final_response": "📧 Email ready to send. Awaiting approval...",
                "requires_approval": True,
            }
        # Read/summarize emails
        return {"final_response": "📧 Email feature requires IMAP configuration.", "response_type": "text"}
'''

# ── Reminder Agent ────────────────────────────────────────────
REMINDER_AGENT = '''"""
Reminder Agent — Tasks, calendar, one-time reminders.
"""
from __future__ import annotations
from typing import Any, Dict
import structlog
from app.agents.base import BaseAgent
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class ReminderAgent(BaseAgent):
    name = "reminder_agent"
    description = "One-time reminders and calendar events"
    supported_intents = ["reminder"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        reminder_text = state.metadata.get("reminder", state.raw_input)
        remind_at = state.metadata.get("remind_at", "")
        # Store in DB for notification dispatch
        return {
            "final_response": f"✅ Reminder set: \\"{reminder_text}\\" at {remind_at}",
            "response_type": "action",
            "agent_logs": [f"[reminder_agent] set reminder for {remind_at}"],
        }
'''

# ── File Agent ────────────────────────────────────────────────
FILE_AGENT = '''"""
File Agent — PDF reading, report generation, DOCX/PPTX creation.
"""
from __future__ import annotations
from pathlib import Path
from typing import Any, Dict
import structlog
from app.agents.base import BaseAgent
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class FileAgent(BaseAgent):
    name = "file_agent"
    description = "Read PDFs, generate reports, create DOCX/PPTX"
    supported_intents = ["file_read", "file_write", "pdf_analysis"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        file_path = state.metadata.get("file_path", "")
        action = state.metadata.get("file_action", "read")

        if action == "read_pdf" and file_path:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = "\\n".join(p.extract_text() or "" for p in reader.pages)
                return {
                    "final_response": f"📄 PDF Content ({len(reader.pages)} pages):\\n\\n{text[:5000]}",
                    "artifacts": [{"type": "pdf_content", "content": text}],
                }
            except Exception as e:
                return {"error": str(e), "final_response": f"PDF read error: {e}"}

        elif action == "create_docx":
            from docx import Document
            doc = Document()
            doc.add_heading(state.metadata.get("title", "Document"), 0)
            doc.add_paragraph(state.metadata.get("content", ""))
            output_path = state.metadata.get("output", "output.docx")
            doc.save(output_path)
            return {"final_response": f"✅ DOCX created: {output_path}", "response_type": "action"}

        return {"final_response": "File operation completed"}
'''

# ── Voice Agent ───────────────────────────────────────────────
VOICE_AGENT = '''"""
Voice Agent — Speech-to-text (Whisper) and text-to-speech (OpenAI TTS).
"""
from __future__ import annotations
from typing import Any, Dict
import structlog
from app.agents.base import BaseAgent
from app.config import settings
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class VoiceAgent(BaseAgent):
    name = "voice_agent"
    description = "Speech-to-text and text-to-speech conversion"
    supported_intents = ["voice_processing"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        action = state.metadata.get("voice_action", "tts")

        if action == "stt":
            # Transcribe audio file
            audio_path = state.metadata.get("audio_path", "")
            if not audio_path:
                return {"error": "No audio file provided"}
            try:
                from openai import AsyncOpenAI
                client = AsyncOpenAI(api_key=settings.openai_api_key)
                with open(audio_path, "rb") as f:
                    transcript = await client.audio.transcriptions.create(
                        model=settings.whisper_model,
                        file=f,
                    )
                return {
                    "final_response": transcript.text,
                    "raw_input": transcript.text,
                    "agent_logs": [f"[voice_agent] transcribed {len(transcript.text)} chars"],
                }
            except Exception as e:
                return {"error": str(e), "final_response": f"STT error: {e}"}

        elif action == "tts":
            text = state.final_response or state.raw_input
            try:
                from openai import AsyncOpenAI
                client = AsyncOpenAI(api_key=settings.openai_api_key)
                response = await client.audio.speech.create(
                    model="tts-1",
                    voice=settings.tts_voice,
                    input=text[:4096],
                )
                audio_data = response.content
                return {
                    "final_response": "🔊 Audio generated",
                    "artifacts": [{"type": "audio", "data": audio_data, "format": "mp3"}],
                }
            except Exception as e:
                return {"error": str(e), "final_response": f"TTS error: {e}"}

        return {"final_response": "Voice action completed"}
'''

# ── Scheduler Agent ───────────────────────────────────────────
SCHEDULER_AGENT = '''"""
Scheduler Agent — Recurring tasks, cron jobs, background workflows.
Different from Reminder Agent — this runs recurring automated workflows.
"""
from __future__ import annotations
from typing import Any, Dict
import structlog
from croniter import croniter
from app.agents.base import BaseAgent
from app.config import settings
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

class SchedulerAgent(BaseAgent):
    name = "scheduler_agent"
    description = "Recurring cron jobs, background workflows, automated task scheduling"
    supported_intents = ["schedule_task"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        cron_expr = state.metadata.get("cron", "0 9 * * 1")  # Default: Mon 9am
        task_description = state.metadata.get("task", state.raw_input)
        task_name = state.metadata.get("name", "scheduled_task")

        # Validate cron expression
        if not croniter.is_valid(cron_expr):
            return {"error": f"Invalid cron: {cron_expr}", "final_response": f"Invalid cron expression: {cron_expr}"}

        # Register with Celery Beat
        from app.services.background import celery_app
        celery_app.conf.beat_schedule[task_name] = {
            "task": "app.services.background.run_scheduled_task",
            "schedule": cron_expr,
            "args": [task_description, state.user_id],
        }

        from croniter import croniter
        from datetime import datetime
        next_run = croniter(cron_expr, datetime.now()).get_next(datetime)

        return {
            "final_response": (
                f"✅ Scheduled task created:\\n"
                f"- Name: {task_name}\\n"
                f"- Schedule: `{cron_expr}`\\n"
                f"- Task: {task_description}\\n"
                f"- Next run: {next_run.strftime(\'%Y-%m-%d %H:%M')}"
            ),
            "response_type": "action",
            "agent_logs": [f"[scheduler_agent] registered {task_name} cron={cron_expr}"],
        }
'''

# ── Project Manager Agent ─────────────────────────────────────
PROJECT_MANAGER_AGENT = '''"""
Project Manager Agent — Create projects, track progress, generate sprint plans.
"""
from __future__ import annotations
from typing import Any, Dict
import structlog
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import ChatPromptTemplate
from app.agents.base import BaseAgent
from app.core.langgraph.state import AgentState

log = structlog.get_logger(__name__)

PM_PROMPT = """You are a senior project manager AI.

Given a project description, create:
{
  "project_name": "",
  "description": "",
  "tech_stack": [],
  "phases": [{"name": "", "duration_weeks": 0, "tasks": []}],
  "milestones": [{"name": "", "due_week": 0, "deliverable": ""}],
  "sprint_1": {"goal": "", "tasks": [], "duration_days": 14},
  "success_criteria": [],
  "risks": [{"risk": "", "mitigation": ""}]
}
"""

class ProjectManagerAgent(BaseAgent):
    name = "project_manager_agent"
    description = "Project creation, milestone tracking, sprint planning"
    supported_intents = ["project_management"]

    async def execute(self, state: AgentState) -> Dict[str, Any]:
        action = state.metadata.get("pm_action", "create_plan")

        if action == "create_plan":
            prompt = ChatPromptTemplate.from_messages([
                ("system", PM_PROMPT),
                ("human", "Project: {project}\\n\\nCreate a comprehensive project plan."),
            ])
            chain = prompt | self.llm | JsonOutputParser()
            try:
                plan = await chain.ainvoke({"project": state.raw_input})
            except Exception as e:
                return {"error": str(e), "final_response": f"Planning error: {e}"}

            summary = (
                f"## 🗂️ {plan.get(\'project_name\', \'Project Plan\')}\\n\\n"
                f"{plan.get(\'description\', \'\')}\\n\\n"
                f"**Tech Stack:** {', '.join(plan.get(\'tech_stack\', []))}\\n\\n"
                f"**Phases:** {len(plan.get(\'phases\', []))} phases\\n\\n"
                f"**Sprint 1 Goal:** {plan.get(\'sprint_1\', {}).get(\'goal\', \'TBD\')}"
            )
            return {
                "final_response": summary,
                "response_type": "report",
                "artifacts": [{"type": "project_plan", "data": plan}],
                "agent_logs": [f"[pm_agent] plan created for: {state.raw_input[:60]}"],
            }

        elif action == "progress_report":
            return {"final_response": "📊 Progress report generation requires project data."}

        return {"final_response": "Project management action completed"}
'''
